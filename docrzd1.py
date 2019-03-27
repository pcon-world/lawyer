# -*- coding: utf-8 -*-
# coding=<utf-8>

import inputcls

def doc1():
	
	from docx import Document
	from docx.shared import Inches, Pt

	document = Document('file.docx')

	section = document.sections[0]
	header = section.header
	paragraph = header.paragraphs[0]
	paragraph.text = ('Created by @lawyer_bot')
	paragraph.alignment = 2

	h1 = document.add_paragraph(u'В ')
	h1.add_run(inputcls.Law.name).bold = True
	h1.add_run(u'\n')
	h1.add_run(inputcls.Law.adress)
	h1.add_run(u'\nИстец:\n').bold = True
	h1.add_run(inputcls.Plf.name)
	h1.add_run(u'\n')
	h1.add_run(inputcls.Plf.adress)
	h1.add_run(u'\nОтветчик:\n').bold = True
	h1.add_run(inputcls.Dft.name)
	h1.add_run(u'\n')
	h1.add_run(inputcls.Dft.ogrn)
	h1.add_run(u'\n')
	h1.add_run(inputcls.Dft.adress)
	h1.add_run(u'\nПредставитель ответчика:\n').bold = True
	h1.add_run(inputcls.Vkl.name)
	h1.add_run(u'\n')
	h1.add_run(inputcls.Vkl.tel)
	h1.add_run(u'\n')
	h1.add_run(inputcls.Vkl.email)
	h1.add_run(u'\nДело № ')
	h1.add_run(inputcls.Lawsuit.number).bold = True
	h1.add_run(u'\n')
	h1.add_run(inputcls.Lawsuit.link)
	h1.paragraph_format.left_indent = Inches(3.5)
	h1.alignment = 0

	title1 = document.add_paragraph(u'Ходатайство')
	title1.add_run(u'\nОб ознакомлении с материалами дела')
	title1.alignment = 1
	
	p1 = document.add_paragraph(u'\tВ производстве')
	p1.add_run(inputcls.Law.name)
	p1.add_run(u'суда находится дело № ')
	p1.add_run(inputcls.Lawsuit.number)
	p1.add_run(u' по иску ')
	p1.add_run(inputcls.Plf.name)
	p1.add_run(u' к ')
	p1.add_run(inputcls.Dft.name)
	p1.add_run(u'.\n\tНа основании ст. 41 АПК РФ')

	p2 = document.add_paragraph(u'ПРОШУ СУД:')
	p2.alignment = 1

	p3 = document.add_paragraph(u'\tПредставить возможность ознакомиться с материалами дела №')
	p3.add_run(inputcls.Lawsuit.number)
	p3.add_run(u' посредством снятия фотокопий с указанных материалов.')

	document.add_paragraph(u'К заявлению прилагаю следующие документы:')
	ft_d1 = document.add_paragraph(
		u'Копия доверенности № ', style='List Number')

	foot_data = document.add_paragraph(u'Представитель по доверенности\n')
	inputcls.Vkl.name = inputcls.Vkl.name.split(' ')
	inputcls.Vkl.name[1] = inputcls.Vkl.name[1][0]
	inputcls.Vkl.name[2] = inputcls.Vkl.name[2][0]
	foot_data.add_run(inputcls.Vkl.name)

	document.add_page_break()

	#document.save('test.docx')
	#print(inputcls.Law.name)
	file_name = inputcls.Law.name+'.docx'

	document.save(file_name)